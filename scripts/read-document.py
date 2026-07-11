# read-document.py
# Doc file .docx va .pdf, xuat ra markdown text
# Usage: python scripts/read-document.py <file-path>

import sys
import os
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')


def read_docx(path):
    from docx import Document
    doc = Document(path)

    lines = []
    for p in doc.paragraphs:
        style = p.style.name.lower() if p.style and p.style.name else ''
        text = p.text.strip()
        if not text:
            lines.append('')
            continue
        if 'heading' in style:
            level = style.replace('heading ', '').replace('headings ', '')
            try:
                n = int(level)
            except ValueError:
                n = 1
            lines.append(f"{'#' * n} {text}")
        else:
            lines.append(text)

    if doc.tables:
        lines.append('')
        for i, table in enumerate(doc.tables):
            if i > 0:
                lines.append('')
            rows = []
            col_widths = []
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                if not col_widths:
                    col_widths = [max(len(c) + 2 for c in cells)]
                else:
                    cw = [max(col_widths[j] if j < len(col_widths) else 0, len(cells[j]) + 2) for j in range(len(cells))]
                    col_widths = cw
                rows.append(cells)

            for ri, row in enumerate(rows):
                line = '| ' + ' | '.join(c.ljust(col_widths[j] if j < len(col_widths) else len(c) + 2) for j, c in enumerate(row)) + ' |'
                lines.append(line)
                if ri == 0:
                    sep = '|-' + '-|-'.join('-' * (col_widths[j] if j < len(col_widths) else 3) for j in range(len(row))) + '-|'
                    lines.append(sep)

    return '\n'.join(lines)


def read_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    lines = []
    for i, page in enumerate(reader.pages):
        if i > 0:
            lines.append('')
            lines.append(f'--- page {i + 1} ---')
            lines.append('')
        text = page.extract_text()
        if text:
            lines.append(text.strip())
    return '\n'.join(lines)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python scripts/read-document.py <file-path>', file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f'ERROR: File not found: {file_path}', file=sys.stderr)
        sys.exit(1)

    ext = Path(file_path).suffix.lower()
    if ext == '.docx':
        print(read_docx(file_path))
    elif ext == '.pdf':
        print(read_pdf(file_path))
    else:
        print(f'ERROR: Unsupported file type: {ext}', file=sys.stderr)
        sys.exit(1)
